# =====================================================
# app/utils/file_processors.py
import os
import uuid
from typing import Tuple, Optional, Dict, Any
import PyPDF2
from docx import Document as DocxDocument
import mammoth
from pathlib import Path

from app.core.config import settings
from app.utils.text_processors import TextCleaner, TextChunker
from app.models.document import DocumentType

class FileProcessor:
    """Process different file types and extract text"""
    
    def __init__(self):
        self.text_cleaner = TextCleaner()
        self.text_chunker = TextChunker()
    
    def process_file(self, file_path: str, original_filename: str) -> Dict[str, Any]:
        """Process uploaded file and extract text"""
        try:
            # Determine file type
            file_extension = Path(original_filename).suffix.lower()
            document_type = self._get_document_type(file_extension)
            
            if not document_type:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Extract text based on file type
            raw_text, metadata = self._extract_text(file_path, document_type)
            
            # Clean the text
            cleaned_text = self.text_cleaner.clean_text(raw_text)
            
            # Generate document ID
            document_id = str(uuid.uuid4())
            
            # Create chunks
            chunks = self.text_chunker.chunk_text(cleaned_text, document_id)
            
            # Calculate statistics
            file_stats = os.stat(file_path)
            
            return {
                "document_id": document_id,
                "original_filename": original_filename,
                "document_type": document_type,
                "file_size": file_stats.st_size,
                "raw_text": raw_text,
                "cleaned_text": cleaned_text,
                "chunks": chunks,
                "metadata": {
                    **metadata,
                    "total_chunks": len(chunks),
                    "word_count": len(cleaned_text.split()),
                    "char_count": len(cleaned_text)
                }
            }
            
        except Exception as e:
            raise Exception(f"Error processing file {original_filename}: {str(e)}")
    
    def _get_document_type(self, file_extension: str) -> Optional[DocumentType]:
        """Determine document type from file extension"""
        extension_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TXT
        }
        return extension_map.get(file_extension)
    
    def _extract_text(self, file_path: str, document_type: DocumentType) -> Tuple[str, Dict[str, Any]]:
        """Extract text from file based on document type"""
        metadata = {}
        
        if document_type == DocumentType.PDF:
            return self._extract_from_pdf(file_path, metadata)
        elif document_type == DocumentType.DOCX:
            return self._extract_from_docx(file_path, metadata)
        elif document_type == DocumentType.TXT:
            return self._extract_from_txt(file_path, metadata)
        else:
            raise ValueError(f"Unsupported document type: {document_type}")
    
    def _extract_from_pdf(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text = ""
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += f"\n[Page {page_num + 1}]\n{page_text}\n"
                    except Exception as e:
                        print(f"Error extracting page {page_num + 1}: {e}")
                        continue
                
                metadata["page_count"] = page_count
                return text, metadata
                
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from DOCX file"""
        try:
            # Method 1: Using python-docx (preserves structure)
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # If that didn't work well, try mammoth (better formatting)
            if len(text.strip()) < 100:  # Fallback to mammoth
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
            
            metadata["paragraph_count"] = len(doc.paragraphs)
            return text, metadata
            
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path: str, metadata: Dict) -> Tuple[str, Dict]:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    metadata["encoding"] = encoding
                    return text, metadata
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with any common encoding")
            
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
